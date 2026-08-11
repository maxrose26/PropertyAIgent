"""Document download + text extraction + priority-ordered combination.

Ported from extract_pdf_text.py (pdfplumber) and
build_document_extraction_notes.py (document-type classification and the
priority-ordered "combined_priority_text" that feeds the LLM stages).
"""
from __future__ import annotations

import hashlib
import multiprocessing
import re
from pathlib import Path
from urllib.parse import urlparse

import docx
import pdfplumber
import requests

from app.db.session import DATA_DIR
from app.scrapers.ssl_fix import get_verify_bundle_for_host

HEADERS = {"User-Agent": "Mozilla/5.0"}

PRIORITY_ORDER = [
    "application_form",
    "planning_statement",
    "decision_notice",
    "officer_report",
    "viability_affordable_housing",
    "design_access",
    "s106",
]

# The only document types the extraction pipeline actually reads (see
# build_combined_priority_text below) - everything else (site plans, CAD
# drawings, ecology/drainage/transport reports, landscaping plans...) was
# being downloaded and text-extracted anyway despite never being used,
# confirmed wasteful on a real 248-unit scheme with 150+ documents. Callers
# doing the discovery/download step should skip anything not in this set.
USEFUL_DOC_TYPES = set(PRIORITY_ORDER)

# \s* alone only matches whitespace - fine for live-scraped document names,
# which come from the portal's own display title ("Planning Statement.pdf"),
# but recovered/imported documents can carry filename-derived names instead
# ("DC_086371-PLANNING_STATEMENT-1973058.pdf"), where words are joined by
# underscores or hyphens rather than spaces. Confirmed a real case: an
# entire Planning Statement, Decision Notice and other real documents sat
# unclassified as "other" (silently excluded from extraction) purely
# because "planning_statement" didn't match `planning\s*statement`. SEP
# matches any of the separator styles actually seen across both sources.
SEP = r"[\s_.-]*"

DOCUMENT_TYPE_PATTERNS = [
    ("application_form", [rf"application{SEP}form", rf"app(?:lication)?{SEP}forms?"]),
    ("planning_statement", [rf"planning{SEP}statement", rf"supporting{SEP}statement"]),
    ("decision_notice", [rf"decision{SEP}(notice|letter)", rf"approval{SEP}notice", rf"refusal{SEP}notice"]),
    ("officer_report", [rf"officer{SEP}report", rf"delegated{SEP}report", rf"committee{SEP}report"]),
    ("viability_affordable_housing", [r"viability", rf"affordable{SEP}housing", rf"financial{SEP}viability"]),
    ("design_access", [rf"design{SEP}(and|&){SEP}access", rf"design{SEP}statement"]),
    ("s106", [r"s106", rf"section{SEP}106", rf"legal{SEP}agreement"]),
]

MAX_CHARS_PER_DOC = 60000
MAX_TOTAL_CHARS = 120000


def standardise_document_type(document_name: str, doc_type_raw: str) -> str:
    raw = f"{doc_type_raw} {document_name}".lower()
    for standard_type, patterns in DOCUMENT_TYPE_PATTERNS:
        if any(re.search(pattern, raw, flags=re.I) for pattern in patterns):
            return standard_type
    return "other"


# Bare numeric/generic filenames carry zero classification signal by name
# (confirmed real case: Manchester's Arcus portal stores the council's own
# internal document ID as the "name" for a large share of uploads -
# "805570.pdf", "00868.xlsm" - not the real document title). Deliberately
# narrow: only genuinely uninformative names trigger the content-sniffing
# fallback below, NOT every "other"-classified document - a document
# correctly named "Tree Survey.pdf" or "Site Location Plan.pdf" really is
# a type extraction never reads, and downloading/scanning it anyway would
# be pure waste repeated across every application on every council.
_UNINFORMATIVE_NAME_RE = re.compile(r"^(?:\d+|document\d*|untitled\d*|scan\d*|file\d*|attachment\d*)$", re.I)


def name_is_uninformative(document_name: str) -> bool:
    stem = Path(document_name or "").stem.strip()
    return bool(_UNINFORMATIVE_NAME_RE.match(stem))


def sniff_document_type_from_text(text: str) -> str:
    """Same pattern set as standardise_document_type, applied to the
    document's own extracted text instead of its filename - the fallback for
    name_is_uninformative cases, where the filename gives no signal but the
    document itself still opens with a real title (a Planning Statement's
    first page almost always states "Planning Statement" near the top,
    regardless of what the uploaded file happened to be named). Checked
    against only the first slice of text, matching where a document's own
    title/header actually appears - a long document could otherwise pick up
    an unrelated false match deep in its body (e.g. a Planning Statement's
    "Design and Access" bullet in an unrelated contents-page listing)."""
    head = (text or "")[:3000].lower()
    for standard_type, patterns in DOCUMENT_TYPE_PATTERNS:
        if any(re.search(pattern, head, flags=re.I) for pattern in patterns):
            return standard_type
    return "other"


def _priority_rank(doc_type: str) -> int:
    try:
        return PRIORITY_ORDER.index(doc_type)
    except ValueError:
        return len(PRIORITY_ORDER) + 1


def sanitise_filename(name: str) -> str:
    name = re.sub(r"[^\w\-. ]", "_", name).strip()
    return name[:150] or "document"


def document_dir(council_code: str, reference: str) -> Path:
    folder = DATA_DIR / "documents" / council_code / re.sub(r"[^\w\-]", "_", reference)
    folder.mkdir(parents=True, exist_ok=True)
    return folder


# Bounded chunk size for streamed downloads - just large enough to be
# efficient, small enough that "how much of one file is in memory at once"
# is never a meaningful contributor to process-tree RSS.
DOWNLOAD_CHUNK_SIZE = 256 * 1024  # 256KB

# A DEFENSIVE ceiling, distinct from MAX_EXTRACTABLE_FILE_SIZE below (which
# governs whether a downloaded file is worth text-EXTRACTING, not whether
# it's downloaded at all - the file itself is still kept as evidence per
# this project's own principle of never losing source information, even
# when it's too large to extract). Protects only against a single
# pathological/misbehaving response (a mislabeled non-PDF asset, a server
# that never closes the connection) - comfortably above any genuine
# planning document confirmed seen in production.
MAX_DOWNLOAD_FILE_SIZE = 200 * 1024 * 1024  # 200MB


def download_document(
    council_code: str, reference: str, document_name: str, source_url: str,
    session: requests.Session | None = None, referer: str | None = None,
) -> Path | None:
    folder = document_dir(council_code, reference)
    # document_name alone isn't guaranteed unique within an application - Idox
    # portals commonly reuse names like "Location Plan" across multiple actual
    # files, and a mis-parsed name column can make EVERY document in an
    # application collide onto the same filename (confirmed: this silently
    # destroyed 134 of 135 real downloads for one Oldham application, since
    # the existence check below just returned the first document's file for
    # every subsequent one). A short hash of the source URL - which Idox
    # always makes unique per file - makes collisions structurally impossible
    # regardless of what the name column says.
    url_hash = hashlib.sha1(source_url.encode("utf-8")).hexdigest()[:10]
    dest = folder / f"{url_hash}_{sanitise_filename(document_name)}"

    if dest.exists():
        return dest

    headers = dict(HEADERS)
    if referer:
        headers["Referer"] = referer

    # Some Idox instances (Oldham confirmed) 404 a direct file request unless
    # it carries both a Referer AND the session cookies from having visited
    # the document listing page first - a standalone requests.get() with only
    # the Referer header still 404s. Reuse the caller's session when given.
    getter = session.get if session is not None else requests.get
    # verify=<bundle> rather than requests' default - confirmed a real case
    # (Manchester's Arcus portal) where every single document download
    # failed with SSLCertVerificationError because the server's own TLS
    # handshake omits its intermediate certificate (see ssl_fix's module
    # docstring). A no-op for any normally-configured host.
    verify = get_verify_bundle_for_host(urlparse(source_url).hostname)
    # stream=True (Render Daily Discovery Salford document-stage memory
    # diagnosis, Part 8): previously a plain non-streaming GET, which
    # requests fully buffers the ENTIRE response body into response.content
    # before this function gets any chance to look at its size - a single
    # large real-world planning document (confirmed production cases:
    # multi-hundred-MB Design & Access Statements with embedded drawings)
    # would sit ENTIRELY in RAM here, regardless of the separate 15MB
    # MAX_EXTRACTABLE_FILE_SIZE cap below, which only ever gets consulted
    # AFTER this function already returns - too late to prevent the spike.
    # Streaming + writing straight to disk in bounded chunks means at most
    # one DOWNLOAD_CHUNK_SIZE's worth of one document is ever held in memory
    # at a time, independent of how large the real file turns out to be.
    response = getter(source_url, headers=headers, timeout=60, verify=verify, stream=True)
    try:
        response.raise_for_status()
        downloaded = 0
        tmp_dest = dest.with_name(dest.name + ".part")
        try:
            with tmp_dest.open("wb") as f:
                for chunk in response.iter_content(chunk_size=DOWNLOAD_CHUNK_SIZE):
                    if not chunk:
                        continue
                    downloaded += len(chunk)
                    if downloaded > MAX_DOWNLOAD_FILE_SIZE:
                        raise ValueError(
                            f"{source_url} exceeded {MAX_DOWNLOAD_FILE_SIZE} bytes while "
                            "downloading - aborted to protect process memory"
                        )
                    f.write(chunk)
            tmp_dest.replace(dest)
        except BaseException:
            tmp_dest.unlink(missing_ok=True)
            raise
    finally:
        response.close()
    return dest


def extract_pdf_text(path: Path) -> str:
    parts: list[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
    except Exception:
        return ""
    return "\n".join(parts)


def extract_docx_text(path: Path) -> str:
    """Planning documents (viability statements especially) are frequently
    lodged as .docx rather than PDF - without this, those documents were
    downloaded but silently never read."""
    try:
        document = docx.Document(path)
        parts = [p.text for p in document.paragraphs if p.text]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    except Exception:
        return ""


def sniff_file_kind(path: Path) -> str:
    """The council portal's own file-type label (used for the download's
    filename extension) isn't always right - confirmed a document catalogued
    as "Viability statement.docx" was actually a PDF. Detect by magic bytes
    instead of trusting the extension."""
    try:
        header = path.open("rb").read(8)
    except Exception:
        return "unknown"
    if header.startswith(b"%PDF"):
        return "pdf"
    if header.startswith(b"PK\x03\x04"):
        return "docx"  # zip-based - docx, xlsx etc. all share this signature
    return "unknown"


# Large files (78MB Design & Access Statement, almost entirely embedded
# architectural drawings) confirmed sending pdfplumber's memory past 9GB.
# Cheap first filter - skip the obviously-huge ones without even opening them.
MAX_EXTRACTABLE_FILE_SIZE = 15 * 1024 * 1024  # 15MB

# File size alone isn't a reliable proxy for parsing cost, though - confirmed
# a SECOND real case where a file under this size limit *still* blew past
# 7GB, almost certainly a densely vectorised CAD-style drawing (highly
# compressed on disk, enormous once pdfplumber catalogues every curve/path
# as an object in memory). The only fix that actually bounds this is a hard,
# forcible timeout - run extraction in its own subprocess so a pathological
# document can be killed outright without threatening the rest of the run
# (a thread-based timeout can't do this: Python can't force-stop a thread
# stuck inside a C/pure-Python parsing loop, only stop *waiting* on it).
EXTRACTION_TIMEOUT_SECONDS = 60


def _extract_in_subprocess(path_str: str, kind: str, result_queue) -> None:
    try:
        path = Path(path_str)
        text = extract_pdf_text(path) if kind == "pdf" else extract_docx_text(path)
        result_queue.put(text)
    except Exception:
        result_queue.put("")


def extract_document_text(path: Path) -> str:
    try:
        if path.stat().st_size > MAX_EXTRACTABLE_FILE_SIZE:
            return ""
    except OSError:
        return ""
    kind = sniff_file_kind(path)
    if kind not in ("pdf", "docx"):
        return ""  # legacy .doc, images, CAD drawings etc. - not text-extractable here

    ctx = multiprocessing.get_context("spawn")
    result_queue = ctx.Queue()
    process = ctx.Process(target=_extract_in_subprocess, args=(str(path), kind, result_queue))
    process.start()
    # Read from the queue BEFORE joining, not after - confirmed a real
    # deadlock: multiprocessing.Queue.put() hands off to a background feeder
    # thread that writes through an OS pipe with limited buffer capacity: a
    # large result (a 422K-char extraction) doesn't fit, so the feeder
    # thread blocks until something reads the pipe - but the child can't
    # fully exit until that thread finishes, and process.join() called
    # first just blocks waiting for exit without ever reading. The result:
    # a document that finished its real work in 17.8s and called put()
    # still showed process.is_alive()==True after the entire timeout, every
    # time, because parent and child were each waiting on the other.
    try:
        text = result_queue.get(timeout=EXTRACTION_TIMEOUT_SECONDS)
    except Exception:
        text = ""
    if process.is_alive():
        process.terminate()
        process.join(timeout=5)
        if process.is_alive():
            process.kill()
    else:
        process.join(timeout=5)
    return text


def clean_document_text(text: str, max_chars: int = MAX_CHARS_PER_DOC) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[\t\r\f\v]+", " ", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    text = text.strip()
    if max_chars and len(text) > max_chars:
        text = text[:max_chars].rstrip() + "\n[TRUNCATED]"
    return text


def build_combined_priority_text(
    docs: list[tuple[str, str, str]], max_total_chars: int = MAX_TOTAL_CHARS
) -> str:
    """docs: list of (doc_type, document_name, clean_text), priority-ordered
    and length-truncated the same way the prototype's build_sections() did."""
    ordered = sorted(docs, key=lambda d: (_priority_rank(d[0]), -len(d[2])))

    parts: list[str] = []
    total = 0
    for doc_type, doc_name, text in ordered:
        if not text:
            continue
        header = f"=== {doc_type.replace('_', ' ').upper()} ===\nDocument: {doc_name}\n"
        section = f"{header}\n{text}\n"

        if max_total_chars and total + len(section) > max_total_chars:
            remaining = max_total_chars - total
            if remaining <= 500:
                break
            parts.append(section[:remaining].rstrip() + "\n[COMBINED TEXT TRUNCATED]")
            break

        parts.append(section)
        total += len(section)

    return "\n\n".join(parts).strip()


def build_broad_sample_text(
    docs: list[tuple[str, str, str]], max_chars_per_doc: int = 2500, max_total_chars: int = 40000
) -> str:
    """Priority-ordered + depth-first (build_combined_priority_text) is right
    for understanding *one* document deeply, but company/person names can
    turn up in any document - confirmed a developer name sitting only in the
    viability statement, which a 60k-char planning statement + application
    form combo pushed completely outside a 22k-char extraction window. For
    entity-hunting, sample a small slice of *every* document instead of a
    large slice of the top few."""
    ordered = sorted(docs, key=lambda d: _priority_rank(d[0]))

    parts: list[str] = []
    total = 0
    for doc_type, doc_name, text in ordered:
        if not text:
            continue
        snippet = text[:max_chars_per_doc]
        header = f"=== {doc_type.replace('_', ' ').upper()} ===\nDocument: {doc_name}\n"
        section = f"{header}\n{snippet}\n"

        if max_total_chars and total + len(section) > max_total_chars:
            break

        parts.append(section)
        total += len(section)

    return "\n\n".join(parts).strip()
